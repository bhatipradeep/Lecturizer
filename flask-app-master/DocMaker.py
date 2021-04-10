import docx
import os


class DocMaker:

    def __init__(self,file_name):
        self.doc = docx.Document()
        self.file_name = file_name

    def save(self):
        if not os.path.exists('C:/lecturizer_reports'):
            os.makedirs('C:/lecturizer_reports')

        self.doc.save("C:/lecturizer_reports/"+self.file_name+".docx")

    def add_summary(self,text):
        self.doc.add_heading("SUMMARY",2)
        self.doc.add_paragraph(text)
        self.save()
    
    def add_transcription(self,text):
        self.doc.add_heading("Transcription",2)
        self.doc.add_paragraph(text)
        self.save()

    def add_imp_words(self,words):
        self.doc.add_heading("Important Terms",2)
        for word in words:
            self.doc.add_paragraph(word)
        self.save()
    

